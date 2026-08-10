#!/usr/bin/env python3
"""md_to_docx.py — Convert a Markdown file to a Microsoft Word (.docx) document.

Depends only on `python-docx` (no pandoc required).

Supports:
  - ATX headings (# .. ######) with inline formatting
  - Paragraphs (soft-wrapped lines joined)
  - Bold (**), italic (*), inline code (`), strikethrough (~~), links [t](u)
  - Pipe tables with per-column alignment (left / right / center)
  - Blockquotes (rendered as indented callout paragraphs)
  - Bullet, ordered, nested, and task / checkbox lists
  - Horizontal rules
  - Optional YAML frontmatter (title / author / date parsed, body skipped)

Usage:
    python3 md_to_docx.py -i report.md -o report.docx --title "Report" --author "Me"
"""
import argparse
import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# --------------------------------------------------------------------------
# Inline parsing — token order matters: code, bold, strike, link, italic
# --------------------------------------------------------------------------
INLINE_RE = re.compile(r"(`[^`]+`|\*\*.+?\*\*|~~.+?~~|\[.+?\]\(.+?\)|\*.+?\*)")
LINK_RE = re.compile(r"^\[(.+?)\]\((.+?)\)$")
ALIGN = {"left": WD_ALIGN_PARAGRAPH.LEFT, "right": WD_ALIGN_PARAGRAPH.RIGHT,
         "center": WD_ALIGN_PARAGRAPH.CENTER}


def add_run(p, text, bold=False, italic=False, strike=False, code=False):
    r = p.add_run(text)
    if bold:
        r.bold = True
    if italic:
        r.italic = True
    if strike:
        r.font.strike = True
    if code:
        r.font.name = "Consolas"
        r.font.size = Pt(10)
    return r


def add_hyperlink(p, url, text):
    """Append a clickable hyperlink to paragraph p (appends at end, in order)."""
    r_id = p.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color = OxmlElement("w:color"); color.set(qn("w:val"), "0563C1"); rPr.append(color)
    underline = OxmlElement("w:u"); underline.set(qn("w:val"), "single"); rPr.append(underline)
    run.append(rPr)
    t = OxmlElement("w:t"); t.text = text; run.append(t)
    hyperlink.append(run)
    p._p.append(hyperlink)


def add_inline(p, text):
    """Parse inline markdown in `text` and append formatted runs to paragraph p."""
    pos = 0
    for m in INLINE_RE.finditer(text):
        if m.start() > pos:
            add_run(p, text[pos:m.start()])
        tok = m.group(0)
        if tok.startswith("`"):
            add_run(p, tok[1:-1], code=True)
        elif tok.startswith("**"):
            add_run(p, tok[2:-2], bold=True)
        elif tok.startswith("~~"):
            add_run(p, tok[2:-2], strike=True)
        elif tok.startswith("["):
            lm = LINK_RE.match(tok)
            if lm:
                add_hyperlink(p, lm.group(2), lm.group(1))
            else:
                add_run(p, tok)
        elif tok.startswith("*"):
            add_run(p, tok[1:-1], italic=True)
        pos = m.end()
    if pos < len(text):
        add_run(p, text[pos:])


# --------------------------------------------------------------------------
# Block detection
# --------------------------------------------------------------------------
BULLET_RE = re.compile(r"^(\s*)([-*+])\s+(.*)$")
ORDERED_RE = re.compile(r"^(\s*)(\d+[.)])\s+(.*)$")
# A line that begins a new paragraph within a blockquote (list/clause markers)
MARKER_RE = re.compile(
    r"^\s*(?:[-*+]\s+|\d+[.)]\s+|\([a-z0-9]{1,3}\)\s+|[ivxlcdm]+[.)]\s+)")


def is_table(line):     return line.strip().startswith("|")
def is_blockquote(line):return line.strip().startswith(">")
def is_list(line):      return bool(BULLET_RE.match(line)) or bool(ORDERED_RE.match(line))
def is_heading(line):
    return bool(re.match(r"^#{1,6}\s+\S", line.strip()))


def is_hr(line):
    s = line.strip()
    return s in ("---", "***", "___") or bool(re.fullmatch(r"[-*_]{3,}", s))


def block_starts(line):
    s = line.strip()
    return (s == "" or is_heading(line) or is_table(line) or is_blockquote(line)
            or is_list(line) or is_hr(line))


# --------------------------------------------------------------------------
# Builders
# --------------------------------------------------------------------------
def add_hr(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1"); bottom.set(qn("w:color"), "auto")
    pbdr.append(bottom); pPr.append(pbdr)


def add_heading(doc, text, level):
    p = doc.add_heading("", level=min(level, 9))
    add_inline(p, text)


def add_paragraph(doc, text):
    p = doc.add_paragraph()
    add_inline(p, text)
    return p


def add_blockquote(doc, bq_lines):
    """Render blockquote lines as indented callout paragraphs.

    Soft-wrapped lines join into one paragraph; a blank `>` line or a line that
    begins a list/clause marker starts a new paragraph.
    """
    contents = []
    for raw in bq_lines:
        s = raw.strip()
        if s.startswith(">"):
            s = s[1:].lstrip()
        contents.append(s)

    groups, para = [], []
    for s in contents:
        if s == "":
            if para:
                groups.append(para); para = []
            continue
        if para and MARKER_RE.match(s):
            groups.append(para); para = []
        para.append(s)
    if para:
        groups.append(para)

    for g in groups:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(4)
        add_inline(p, " ".join(x.strip() for x in g))


def add_list(doc, items):
    bullets = ["\u2022", "\u25e6", "\u25aa"]  # •  ◦  ▪
    for raw in items:
        mb, mo = BULLET_RE.match(raw), ORDERED_RE.match(raw)
        if mb:
            indent_spaces, content, ordered, marker = len(mb.group(1)), mb.group(3), False, mb.group(2)
        elif mo:
            indent_spaces, content, ordered, marker = len(mo.group(1)), mo.group(3), True, mo.group(2)
        else:
            indent_spaces, content, ordered, marker = 0, raw.strip(), False, "-"

        level = indent_spaces // 2
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3 + 0.25 * level)
        p.paragraph_format.space_after = Pt(2)

        cb = re.match(r"\[( |x|X)\]\s+(.*)", content)
        if cb:
            p.add_run("\u2612  " if cb.group(1).lower() == "x" else "\u2610  ")
            add_inline(p, cb.group(2))
        elif ordered:
            p.add_run(f"{marker}  ")
            add_inline(p, content)
        else:
            p.add_run(f"{bullets[min(level, 2)]}  ")
            add_inline(p, content)


def parse_row(line):
    s = line.strip()
    if s.startswith("|"):
        s = s[1:]
    if s.endswith("|"):
        s = s[:-1]
    return [c.strip() for c in s.split("|")]


def col_align(cell):
    c = cell.strip()
    if c.startswith(":") and c.endswith(":"):
        return "center"
    if c.endswith(":"):
        return "right"
    return "left"


def is_sep_row(line):
    cells = parse_row(line)
    return bool(cells) and all(re.fullmatch(r":?-{1,}:?", c.strip()) for c in cells)


def set_cell(cell, text, bold=False, align="left"):
    p = cell.paragraphs[0]
    for r in list(p.runs):
        r._r.getparent().remove(r._r)
    add_inline(p, text)
    if bold:
        for r in p.runs:
            r.bold = True
    p.alignment = ALIGN.get(align, WD_ALIGN_PARAGRAPH.LEFT)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)


def add_table(doc, table_lines):
    rows = [parse_row(l) for l in table_lines]
    if not rows:
        return
    header = rows[0]
    aligns = ["left"] * len(header)
    data = rows[1:]
    if len(rows) > 1 and is_sep_row(table_lines[1]):
        aligns = [col_align(c) for c in rows[1]]
        data = rows[2:]
    n_cols = max(len(r) for r in ([header] + data))
    table = doc.add_table(rows=1, cols=n_cols)
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    for j, h in enumerate(header):
        a = aligns[j] if j < len(aligns) else "left"
        set_cell(table.rows[0].cells[j], h, bold=True, align=a)
    for d in data:
        cells = table.add_row().cells
        for j in range(n_cols):
            a = aligns[j] if j < len(aligns) else "left"
            set_cell(cells[j], d[j] if j < len(d) else "", align=a)


# --------------------------------------------------------------------------
# Frontmatter + main conversion
# --------------------------------------------------------------------------
def parse_frontmatter(lines):
    fm = {}
    if not lines or lines[0].strip() != "---":
        return fm, lines
    end = None
    for k in range(1, len(lines)):
        if lines[k].strip() in ("---", "..."):
            end = k
            break
    if end is None:
        return fm, lines
    for line in lines[1:end]:
        if ":" in line:
            key, _, val = line.partition(":")
            fm[key.strip().lower()] = val.strip().strip('"').strip("'")
    return fm, lines[end + 1:]


def convert(md_text, out_path, title=None, author=None, font="Calibri", font_size=11):
    lines = md_text.split("\n")
    fm, lines = parse_frontmatter(lines)
    title = title or fm.get("title")
    author = author or fm.get("author")

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = font
    normal.font.size = Pt(font_size)

    if title:
        doc.core_properties.title = title
    if author:
        doc.core_properties.author = author

    i, n = 0, len(lines)
    while i < n:
        line = lines[i]
        s = line.strip()

        if s == "":
            i += 1
            continue
        if is_hr(line):
            add_hr(doc); i += 1; continue
        if is_heading(line):
            level = len(s) - len(s.lstrip("#"))
            text = s.lstrip("#").rstrip("#").strip()
            add_heading(doc, text, level); i += 1; continue
        if is_table(line):
            block = []
            while i < n and is_table(lines[i]):
                block.append(lines[i].strip()); i += 1
            add_table(doc, block); continue
        if is_blockquote(line):
            block = []
            while i < n and is_blockquote(lines[i]):
                block.append(lines[i]); i += 1
            add_blockquote(doc, block); continue
        if is_list(line):
            block = []
            while i < n and is_list(lines[i]):
                block.append(lines[i]); i += 1
            add_list(doc, block); continue
        # paragraph: join soft-wrapped lines until a block boundary
        buf = []
        while i < n and not block_starts(lines[i]):
            if lines[i].strip():
                buf.append(lines[i].strip())
            i += 1
        if buf:
            add_paragraph(doc, " ".join(buf))

    doc.save(out_path)
    return doc


def main():
    ap = argparse.ArgumentParser(description="Convert Markdown to Word .docx (python-docx).")
    ap.add_argument("-i", "--input", required=True, help="Source Markdown file")
    ap.add_argument("-o", "--output", help="Output .docx path (default: <input>.docx)")
    ap.add_argument("--title", help="Document title (core property); auto from frontmatter if omitted")
    ap.add_argument("--author", help="Document author (core property); auto from frontmatter if omitted")
    ap.add_argument("--font", default="Calibri", help="Body font (default: Calibri)")
    ap.add_argument("--font-size", type=int, default=11, help="Body font size in pt (default: 11)")
    args = ap.parse_args()

    inp = Path(args.input)
    if not inp.exists():
        sys.exit(f"Input not found: {inp}")
    out = Path(args.output) if args.output else inp.with_suffix(".docx")

    convert(inp.read_text(encoding="utf-8"), str(out),
            title=args.title, author=args.author, font=args.font, font_size=args.font_size)

    d = Document(str(out))
    print(f"Wrote {out}")
    print(f"Paragraphs: {len(d.paragraphs)}, Tables: {len(d.tables)}")


if __name__ == "__main__":
    main()
