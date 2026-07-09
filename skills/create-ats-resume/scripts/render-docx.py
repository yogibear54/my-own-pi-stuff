#!/usr/bin/env python3
"""
HTML → DOCX converter for ATS resumes.

Parses the skill's HTML output (single-column, plain HTML as specified in
references/html-structure.md) and writes a clean Microsoft Word document
via python-docx.

Usage:
    python3 render-docx.py <file.html> [more.html ...]

Each input HTML is converted to a DOCX with the same basename next to it.
"""

import sys
import os
from html.parser import HTMLParser

from docx import Document
from docx.shared import Pt, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# A handful of named entities we encounter in resume copy.
ENTITIES = {
    "nbsp": "\u00a0",
    "amp": "&",
    "lt": "<",
    "gt": ">",
    "quot": '"',
    "apos": "'",
    "mdash": "\u2014",
    "ndash": "\u2013",
    "hellip": "\u2026",
    "middot": "\u00b7",
    "bull": "\u2022",
    "times": "\u00d7",
    "reg": "\u00ae",
    "copy": "\u00a9",
    # Greek letters appear in some role titles; add as needed
    "Beta": "\u03b2",
    "alpha": "\u03b1",
    # Common numeric entities
}


class HTMLToDocx(HTMLParser):
    """Minimal HTML parser that emits docx elements.

    Handles the limited HTML vocabulary the resume templates use:
    - block-level: h1/h2/h3, p, ul/li, div (treated as paragraph container),
                  br (treated as line break inside current paragraph)
    - inline: strong/b, em/i, span (no special handling — content passes through)
    - completely ignored: style, script, head, meta, title, link

    Whitespace between tags is collapsed to single spaces; leading/trailing
    spaces per paragraph are trimmed.
    """

    HEADING_LEVEL = {"h1": 0, "h2": 1, "h3": 2}

    def __init__(self, document=None):
        super().__init__(convert_charrefs=True)
        self.doc = document or Document()
        # Set up default document styles
        self._setup_styles()
        # State
        self.paragraph = None  # current docx paragraph being filled
        self.bold = False
        self.italic = False
        self.skip_depth = 0  # >0 means we ignore content (inside <style>/<script>/etc.)
        self.run_buffer = ""  # accumulated text for the next run

    # ── One-time doc setup ──────────────────────────────────────────────

    def _setup_styles(self):
        # Tighten default Normal style to match resume density
        normal = self.doc.styles["Normal"]
        normal.font.name = "Calibri"
        normal.font.size = Pt(11)
        pf = normal.paragraph_format
        pf.space_after = Pt(2)
        pf.space_before = Pt(0)
        # Page margins (~0.75")
        for section in self.doc.sections:
            section.top_margin = Mm(19)
            section.bottom_margin = Mm(19)
            section.left_margin = Mm(19)
            section.right_margin = Mm(19)

    # ── Internal helpers ───────────────────────────────────────────────

    def _flush_run(self):
        """Push accumulated run text into the current paragraph as a run."""
        if not self.run_buffer or self.paragraph is None:
            self.run_buffer = ""
            return
        text = self.run_buffer
        run = self.paragraph.add_run(text)
        if self.bold:
            run.bold = True
        if self.italic:
            run.italic = True
        self.run_buffer = ""

    def _ensure_paragraph(self):
        if self.paragraph is None:
            self.paragraph = self.doc.add_paragraph()

    def _set_normal_paragraph(self):
        """Open a new plain paragraph (after a block tag ends)."""
        self.paragraph = self.doc.add_paragraph()

    def _new_paragraph(self, style_name=None):
        """Open a new paragraph (optionally with a paragraph style)."""
        self._flush_run()
        if style_name:
            try:
                self.paragraph = self.doc.add_paragraph(style=style_name)
            except KeyError:
                self.paragraph = self.doc.add_paragraph()
        else:
            self.paragraph = self.doc.add_paragraph()
        self.run_buffer = ""

    # ── HTMLParser callbacks ───────────────────────────────────────────

    def handle_starttag(self, tag, attrs):
        if tag in ("style", "script", "head", "meta", "link", "title"):
            self.skip_depth += 1
            return

        if self.skip_depth:
            return

        if tag in self.HEADING_LEVEL:
            self._new_paragraph()
            # python-docx heading style; we then keep filling it as text
            # The first run we add will be applied to the heading paragraph.
            self.in_heading = tag
            return

        if tag in ("p", "div"):
            self._new_paragraph()
            return

        if tag == "ul":
            # Bulleted list — handled via List Bullet style on each <li>.
            self._flush_run()
            self.paragraph = None
            return

        if tag == "li":
            self._flush_run()
            self._new_paragraph(style_name="List Bullet")
            return

        if tag in ("br",):
            self._flush_run()
            if self.paragraph is None:
                self._ensure_paragraph()
            self.paragraph.add_run().add_break()
            return

        if tag in ("strong", "b"):
            self._flush_run()
            self.bold = True
            return

        if tag in ("em", "i"):
            self._flush_run()
            self.italic = True
            return

        if tag in ("span",):
            # No structural meaning in ATS HTML; pass content through.
            return

        # Anything else: ignore (e.g., html, body).

    def handle_endtag(self, tag):
        if tag in ("style", "script", "head", "meta", "link", "title"):
            if self.skip_depth > 0:
                self.skip_depth -= 1
            return

        if self.skip_depth:
            return

        if tag in self.HEADING_LEVEL:
            self._flush_run()
            # Convert current paragraph to heading by re-styling
            self._promote_to_heading(tag)
            self.paragraph = None
            self.in_heading = None
            return

        if tag in ("p", "div", "li"):
            self._flush_run()
            self.paragraph = None
            return

        if tag == "ul":
            self._flush_run()
            self.paragraph = None
            return

        if tag in ("strong", "b"):
            self._flush_run()
            self.bold = False
            return

        if tag in ("em", "i"):
            self._flush_run()
            self.italic = False
            return

    def handle_data(self, data):
        if self.skip_depth:
            return
        # Collapse whitespace inside a run; keep newlines as line breaks
        text = data
        if not text:
            return
        # Replace tabs/newlines with single space; collapse consecutive whitespace.
        cleaned = " ".join(text.split())
        self.run_buffer += cleaned

    def handle_entityref(self, name):
        # Some HTML may still arrive as entity refs even with convert_charrefs=True
        self.run_buffer += ENTITIES.get(name, "&" + name + ";")

    def handle_startendtag(self, tag, attrs):
        # self-closing tags like <br />
        self.handle_starttag(tag, attrs)
        self.handle_endtag(tag)

    # ── Heading promotion ──────────────────────────────────────────────

    def _promote_to_heading(self, tag):
        """Re-style the current paragraph as a Heading 1/2/3.

        python-docx models headings via paragraph styles; rebuilding the
        paragraph as a heading is easier than mutating an existing one.
        """
        level = self.HEADING_LEVEL.get(tag, 1)
        text = "".join(r.text for r in self.paragraph.runs)
        # Remove runs from current paragraph
        for r in list(self.paragraph.runs):
            r.text = ""
        # Re-create as a heading
        heading = self.doc.add_heading("", level=level)
        # Carry over bold/italic state from inside the heading
        run = heading.add_run(text)
        if self.bold:
            run.bold = True
        if self.italic:
            run.italic = True
        # Leave the old empty paragraph in place; remove it to keep the doc tidy
        from docx.oxml.ns import qn
        p = self.paragraph._element
        p.getparent().remove(p)
        self.paragraph = heading


# ─────────────────────────────────────────────────────────────────────


def html_to_docx(html_path: str, docx_path: str):
    with open(html_path, "r", encoding="utf-8") as f:
        html = f.read()

    # Strip the surrounding <!DOCTYPE>, <html>, <head>, <body> wrappers so the
    # parser sees just the body content. Several HTMLParser quirks in Python's
    # stdlib cause paragraph loss when DOCTYPE/head are present.
    import re as _re
    body_match = _re.search(r"<body[^>]*>(.*?)</body>", html, flags=_re.DOTALL)
    if body_match:
        html = body_match.group(1)
    # Also drop any inline <style> block (we don't carry CSS into docx).
    html = _re.sub(r"<style[^>]*>.*?</style>", "", html, flags=_re.DOTALL)

    doc = Document()
    parser = HTMLToDocx(document=doc)
    parser.feed(html)
    parser.close()
    doc.save(docx_path)


def main():
    if len(sys.argv) < 2:
        print("Usage: render-docx.py <file.html> [more.html ...]", file=sys.stderr)
        sys.exit(2)

    results = []
    for src in sys.argv[1:]:
        if not os.path.exists(src):
            print(f"  ✗ {src}: not found", file=sys.stderr)
            results.append((src, None, False))
            continue
        out = src[:-5] + ".docx" if src.endswith(".html") else src + ".docx"
        try:
            html_to_docx(src, out)
            size = os.path.getsize(out)
            print(f"  ✓ {out}  ({size:,} bytes)")
            results.append((src, out, True))
        except Exception as e:
            print(f"  ✗ {src}: {e}", file=sys.stderr)
            results.append((src, None, False))

    if not all(ok for _, _, ok in results):
        sys.exit(1)


if __name__ == "__main__":
    main()
